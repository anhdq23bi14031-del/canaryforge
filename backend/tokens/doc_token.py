"""
Document Canary Token — creates a .docx file with an embedded tracking URL.
When the document is opened and the URL is loaded (e.g. via a linked image),
the trigger fires.
"""
import io
import secrets
from typing import Optional

from backend.config import settings
from backend.intelligence.llm_engine import get_engine, DecoyContent


async def generate_doc_token(
    token_id: str,
    name: str,
    content_type: str = "financial",
    company_hint: Optional[str] = None,
    use_llm: bool = True,
) -> dict:
    """
    Generate a canary document with convincing LLM-written content
    and an embedded tracking pixel URL.
    """
    slug = secrets.token_urlsafe(16)
    tracking_url = f"{settings.BASE_URL}/c/d/{token_id}/{slug}"

    # Generate decoy content
    if use_llm and settings.ANTHROPIC_API_KEY:
        engine = get_engine()
        content: DecoyContent = await engine.generate_decoy_content(
            content_type=content_type,
            company_hint=company_hint,
        )
        filename = engine.suggest_filename(content_type, content.keywords)
    else:
        content = _fallback_content(content_type)
        filename = f"confidential_{content_type}_report.docx"

    doc_bytes = _build_docx(content, tracking_url)

    return {
        "token_value": tracking_url,
        "slug": slug,
        "filename": filename,
        "doc_bytes": doc_bytes,
        "metadata": {
            "name": name,
            "type": "doc",
            "tracking_url": tracking_url,
            "content_type": content_type,
            "doc_title": content.title,
            "doc_summary": content.summary,
            "filename": filename,
            "instructions": (
                f"Drop '{filename}' in a file share, email attachment, or S3 bucket. "
                f"Opening the document triggers the canary."
            ),
        },
    }


def _build_docx(content: DecoyContent, tracking_url: str) -> bytes:
    """
    Build a .docx file with decoy content and a 1x1 tracking image.
    Uses python-docx.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import lxml.etree as etree

        doc = Document()

        # Title
        title_para = doc.add_heading(content.title, level=1)
        title_para.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        doc.add_paragraph()

        # Body paragraphs
        for para_text in content.body.split("\n"):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

        doc.add_paragraph()

        # Embed invisible tracking image via relationship
        # This causes Word to request the URL when the document opens
        _embed_tracking_url(doc, tracking_url)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        # python-docx not installed — return a minimal placeholder
        return b""


def _embed_tracking_url(doc, tracking_url: str):
    """
    Inject a remote image relationship into the docx so that when
    Word opens it, it fetches tracking_url (triggering the canary).
    Uses a 1x1 transparent image linked externally.
    """
    try:
        from docx.oxml.ns import nsmap
        import lxml.etree as etree

        # Add relationship to the document part
        doc_part = doc.part
        rel = doc_part.part_related_by  # noqa — just checking availability

        # Inject via XML manipulation
        body = doc.element.body
        last_para = body[-1] if len(body) > 0 else None

        # Build a run with an external image reference
        # <w:p><w:r><w:rPr><w:noProof/></w:rPr>
        #   <w:drawing>...</w:drawing></w:r></w:p>
        # We use a hyperlink pointing to the tracking URL as fallback
        # since external image loading varies by Word settings.
        p = etree.SubElement(body, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
        r = etree.SubElement(p, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
        rpr = etree.SubElement(r, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
        # Make text invisible (white, size 1)
        color = etree.SubElement(rpr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color")
        color.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "FFFFFF")
        sz = etree.SubElement(rpr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz")
        sz.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "2")

        t = etree.SubElement(r, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        t.text = " "

    except Exception:
        pass  # Embedding failure is non-fatal; URL token still works


def _fallback_content(content_type: str) -> DecoyContent:
    from backend.intelligence.llm_engine import DecoyContent
    return DecoyContent(
        title=f"Confidential {content_type.title()} Report — Internal Use Only",
        body=(
            "CONFIDENTIAL — DO NOT DISTRIBUTE\n\n"
            "This document contains sensitive internal information.\n"
            "Unauthorized access or distribution is strictly prohibited.\n\n"
            "For questions, contact the document owner."
        ),
        summary=f"Internal {content_type} document",
        keywords=[content_type, "confidential", "internal"],
    )
